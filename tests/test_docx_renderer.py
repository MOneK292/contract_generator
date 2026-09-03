"""Tests for DOCX placeholder extraction and rendering."""

from __future__ import annotations

import unittest
import zipfile
from pathlib import Path
from tempfile import TemporaryDirectory
from time import perf_counter

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

from src.services.docx.placeholders import PlaceholderExtractor
from src.services.docx.renderer import DocxRenderer


class DocxRendererTest(unittest.TestCase):
    """DOCX renderer behavior."""

    def test_extracts_placeholders(self) -> None:
        with TemporaryDirectory() as temp_dir:
            template_path = self._create_template(Path(temp_dir))

            placeholders = PlaceholderExtractor().extract(template_path)

            self.assertEqual(
                placeholders,
                {
                    "ФИО",
                    "ИНН",
                    "Город",
                    "Ставка",
                    "Дата выдачи",
                    "Компания",
                    "Номер",
                    "Неизвестно",
                },
            )

    def test_replaces_placeholders_in_docx(self) -> None:
        with TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            template_path = self._create_template(root)
            output_path = root / "output.docx"

            data = {
                "ФИО": "Иванов Иван Иванович",
                "ИНН": "123456789012",
                "Город": "Санкт-Петербург",
                "Ставка": "316,00",
                "Дата выдачи": "15.04.2021",
                "Компания": "Ромашка",
                "Номер": "79990000002",
            }

            render_result = DocxRenderer().render(template_path, data, output_path)

            result = Document(output_path)
            body_text = "\n".join(paragraph.text for paragraph in result.paragraphs)
            table_text = "\n".join(
                text
                for table in result.tables
                for text in self._iter_table_text(table)
            )
            header_text = "\n".join(
                paragraph.text
                for section in result.sections
                for paragraph in section.header.paragraphs
            )
            footer_text = "\n".join(
                paragraph.text
                for section in result.sections
                for paragraph in section.footer.paragraphs
            )

            self.assertTrue(render_result.success)
            self.assertEqual(render_result.unresolved_placeholders, ["Неизвестно"])
            self.assertIn("Сотрудник: Иванов Иван Иванович", body_text)
            self.assertIn(
                "Иванов Иван Иванович работает в компании Ромашка с 15.04.2021",
                body_text,
            )
            self.assertIn("Пунктуация: «Иванов Иван Иванович,»", body_text)
            self.assertIn("(123456789012)", body_text)
            self.assertIn("№79990000002", body_text)
            self.assertIn("Город: Санкт-Петербург", body_text)
            self.assertIn("Повтор: Иванов Иван Иванович", body_text)
            self.assertIn("Неизвестный тег: <Неизвестно>", body_text)
            self.assertIn("Разбитый тег: Иванов Иван Иванович", body_text)
            self.assertIn("ИНН: 123456789012", table_text)
            self.assertIn("Ставка: 316,00", table_text)
            self.assertIn("Вложенная таблица: Ромашка", table_text)
            self.assertIn("Дата: 15.04.2021", header_text)
            self.assertIn("Футер: Иванов Иван Иванович", footer_text)

    def test_preserves_formatting_around_placeholders(self) -> None:
        with TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            template_path = root / "formatting.docx"
            output_path = root / "output.docx"
            document = Document()

            paragraph = document.add_paragraph(style="List Number")
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.left_indent = Inches(0.5)
            bold_run = paragraph.add_run("До ")
            bold_run.bold = True
            placeholder_run = paragraph.add_run("<ФИО>")
            placeholder_run.italic = True
            after_run = paragraph.add_run(" после")
            after_run.underline = True
            after_run.font.size = Pt(14)
            after_run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

            table = document.add_table(rows=1, cols=1)
            table.cell(0, 0).text = "Таблица: <Компания>"
            document.save(template_path)

            DocxRenderer().render(
                template_path,
                {
                    "ФИО": "Иванов Иван Иванович",
                    "Компания": "Ромашка",
                },
                output_path,
            )

            rendered = Document(output_path)
            rendered_paragraph = rendered.paragraphs[0]
            self.assertEqual(rendered_paragraph.style.name, "List Number")
            self.assertEqual(rendered_paragraph.alignment, WD_ALIGN_PARAGRAPH.CENTER)
            self.assertEqual(rendered_paragraph.paragraph_format.left_indent, Inches(0.5))
            self.assertEqual(rendered_paragraph.runs[0].text, "До ")
            self.assertTrue(rendered_paragraph.runs[0].bold)
            self.assertEqual(rendered_paragraph.runs[1].text, "Иванов Иван Иванович")
            self.assertTrue(rendered_paragraph.runs[1].italic)
            self.assertEqual(rendered_paragraph.runs[2].text, " после")
            self.assertTrue(rendered_paragraph.runs[2].underline)
            self.assertEqual(rendered_paragraph.runs[2].font.size, Pt(14))
            self.assertEqual(
                rendered_paragraph.runs[2].font.color.rgb,
                RGBColor(0xC0, 0x00, 0x00),
            )
            self.assertEqual(rendered.tables[0].cell(0, 0).text, "Таблица: Ромашка")

    def test_does_not_mutate_data(self) -> None:
        with TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            template_path = self._create_template(root)
            output_path = root / "output.docx"
            data = {"ФИО": "Иванов Иван Иванович"}
            original = dict(data)

            DocxRenderer().render(template_path, data, output_path)

            self.assertEqual(data, original)

    def test_renders_docm_template(self) -> None:
        with TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            docx_template = self._create_template(root)
            docm_template = root / "template.docm"
            docx_template.replace(docm_template)
            self._set_docm_main_content_type(docm_template)
            output_path = root / "output.docx"

            result = DocxRenderer().render(
                docm_template,
                {"Ð¤Ð˜Ðž": "Ð˜Ð²Ð°Ð½Ð¾Ð² Ð˜Ð²Ð°Ð½ Ð˜Ð²Ð°Ð½Ð¾Ð²Ð¸Ñ‡"},
                output_path,
            )

            self.assertTrue(result.success)
            self.assertTrue(output_path.exists())

    def _set_docm_main_content_type(self, path: Path) -> None:
        replacement = (
            "application/vnd.ms-word.document.macroEnabled.main+xml"
        )
        original = (
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"
        )
        temp_path = path.with_suffix(".tmp")
        with zipfile.ZipFile(path, "r") as source:
            with zipfile.ZipFile(temp_path, "w") as target:
                for item in source.infolist():
                    content = source.read(item.filename)
                    if item.filename == "[Content_Types].xml":
                        content = content.replace(
                            original.encode("utf-8"),
                            replacement.encode("utf-8"),
                        )
                    target.writestr(item, content)
        temp_path.replace(path)

    def _create_template(self, root: Path) -> Path:
        document = Document()
        document.add_paragraph("Сотрудник: <ФИО>")
        document.add_paragraph("Город: <Город>")
        document.add_paragraph(
            "<ФИО> работает в компании <Компания> с <Дата выдачи>"
        )
        document.add_paragraph("Пунктуация: «<ФИО>,»")
        document.add_paragraph("(<ИНН>)")
        document.add_paragraph("№<Номер>")
        document.add_paragraph("Повтор: <ФИО>")
        document.add_paragraph("Неизвестный тег: <Неизвестно>")

        split_paragraph = document.add_paragraph("Разбитый тег: ")
        split_paragraph.add_run("<Ф")
        split_paragraph.add_run("И")
        split_paragraph.add_run("О>")

        table = document.add_table(rows=2, cols=1)
        table.cell(0, 0).text = "ИНН: <ИНН>"
        table.cell(1, 0).text = "Ставка: <Ставка>"
        nested_table = table.cell(1, 0).add_table(rows=1, cols=1)
        nested_table.cell(0, 0).text = "Вложенная таблица: <Компания>"

        section = document.sections[0]
        section.header.paragraphs[0].text = "Дата: <Дата выдачи>"
        section.footer.paragraphs[0].text = "Футер: <ФИО>"

        template_path = root / "template.docx"
        document.save(template_path)
        return template_path

    def _iter_table_text(self, table) -> list[str]:
        texts: list[str] = []
        for row in table.rows:
            for cell in row.cells:
                texts.extend(paragraph.text for paragraph in cell.paragraphs)
                for nested_table in cell.tables:
                    texts.extend(self._iter_table_text(nested_table))
        return texts


class DocxRendererBenchmarkTest(unittest.TestCase):
    """DOCX renderer benchmark."""

    def test_renders_large_template_with_reasonable_runtime(self) -> None:
        with TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            template_path = root / "large-template.docx"
            output_path = root / "large-output.docx"
            document = Document()

            for page in range(10):
                for item in range(10):
                    document.add_paragraph(
                        f"Страница {page + 1}, строка {item + 1}: "
                        "<ФИО> / <Компания> / <Дата выдачи>"
                    )
                if page < 9:
                    document.add_page_break()
            document.save(template_path)

            start = perf_counter()
            result = DocxRenderer().render(
                template_path,
                {
                    "ФИО": "Иванов Иван Иванович",
                    "Компания": "Ромашка",
                    "Дата выдачи": "15.04.2021",
                },
                output_path,
            )
            elapsed = perf_counter() - start

            self.assertTrue(result.success)
            self.assertEqual(result.unresolved_placeholders, [])
            self.assertLess(elapsed, 5.0)


if __name__ == "__main__":
    unittest.main()
