"""Integration tests for the full contract generation pipeline."""

from __future__ import annotations

import json
import subprocess
import time
import unittest
from pathlib import Path
from tempfile import TemporaryDirectory
from types import SimpleNamespace
from unittest.mock import patch

from docx import Document

from src.core.contract_engine import ContractEngine, ContractRequest
from src.core.exceptions import RenderingError, TemplateNotFoundError
from src.core.processor_registry import ProcessorRegistry
from src.parsers.employee_parser import EmployeeParser
from src.processors.date_processor import DateProcessor
from src.processors.fio_processor import FioProcessor
from src.processors.money_processor import MoneyProcessor
from src.processors.placeholder_processor import PlaceholderProcessor
from src.services.cleanup.cleanup_service import CleanupService
from src.services.config.template_catalog import TemplateCatalog
from src.services.docx.placeholders import PlaceholderExtractor
from src.services.docx.renderer import DocxRenderer, RenderResult
from src.services.google.drive import GoogleDriveError
from src.services.pdf.converter import PdfConverter


EMPLOYEE_TEXT = "\n".join(
    [
        "ФИО: Иванов Иван Иванович",
        "ИНН: 123456789012",
        "Ставка: 316",
        "Плата: 1000",
        "Дата выдачи: 15.04.2021",
        "День: 14",
        "Месяц: июля",
        "Год: 2026",
    ]
)


class LocalTemplateCache:
    """Test double for Google cache that returns a local template path."""

    def __init__(self, template_path: Path, *, error: Exception | None = None) -> None:
        self.template_path = template_path
        self.error = error

    def get_template(self, file_id: str) -> Path:
        """Return the local test template path."""
        if self.error is not None:
            raise self.error
        return self.template_path


class ContractPipelineIntegrationTest(unittest.TestCase):
    """Integration coverage for the contract generation pipeline."""

    def test_full_pipeline_generates_docx_pdf_and_cleans_temp_files(self) -> None:
        with TemporaryDirectory() as temp_dir:
            context = self._context(Path(temp_dir))

            with self._mock_libreoffice():
                result = context.engine.generate(context.request)

            self.assertTrue(result.success)
            self.assertIsNotNone(result.output_docx)
            self.assertTrue(result.output_docx.exists())
            self.assertTrue(result.output_pdf.exists())
            self.assertGreater(result.output_pdf.stat().st_size, 0)
            self._assert_pdf_has_pages(result.output_pdf)
            self.assertTrue(context.request.output_docx.exists())
            self.assertEqual(list((Path(temp_dir) / "tmp").glob("*.tmp")), [])
            self.assertEqual(list((Path(temp_dir) / "tmp").glob("*.download")), [])

    def test_generated_docx_has_no_placeholders_and_preserves_document_shape(self) -> None:
        with TemporaryDirectory() as temp_dir:
            context = self._context(Path(temp_dir), with_pdf=False)

            result = context.engine.generate(context.request)

            self.assertTrue(result.success)
            self.assertEqual(result.unresolved_placeholders, [])
            self.assertEqual(PlaceholderExtractor().extract(result.output_docx), set())

            rendered = Document(result.output_docx)
            text = self._document_text(rendered)
            self.assertNotRegex(text, r"<[^<>]+>")
            self.assertIn("Иванов Иван Иванович", text)
            self.assertIn("Иванов", text)
            self.assertIn("316,00", text)
            self.assertIn("1000,00", text)
            self.assertIn("15.04.2021", text)
            self.assertIn("июля", text)
            self.assertIn("14", text)
            self.assertIn("2026", text)
            self.assertGreaterEqual(len(rendered.tables), 1)
            self.assertIn("Header Иванов Иван Иванович", self._headers_text(rendered))
            self.assertIn("Footer 123456789012", self._footers_text(rendered))

    def test_smoke_pipeline_call_produces_pdf_without_exceptions(self) -> None:
        with TemporaryDirectory() as temp_dir:
            context = self._context(Path(temp_dir))

            with self._mock_libreoffice():
                result = context.engine.generate(context.request)

            self.assertTrue(result.success)
            self.assertTrue(result.output_pdf.exists())
            self.assertGreater(result.output_pdf.stat().st_size, 0)

    def test_regression_generates_same_contract_twice(self) -> None:
        with TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            context = self._context(root)
            second_request = ContractRequest(
                project_id="project",
                vacancy_id="developer",
                raw_employee_text=EMPLOYEE_TEXT,
                output_docx=root / "tmp" / "contract-second.docx",
            )

            with self._mock_libreoffice():
                first = context.engine.generate(context.request)
                second = context.engine.generate(second_request)

            self.assertTrue(first.output_pdf.exists())
            self.assertTrue(second.output_pdf.exists())
            self.assertEqual(first.output_pdf.read_bytes(), second.output_pdf.read_bytes())
            self.assertTrue(context.request.output_docx.exists())
            self.assertTrue(second_request.output_docx.exists())
            self.assertEqual(list((root / "tmp").glob("*.tmp")), [])
            self.assertEqual(list((root / "tmp").glob("*.download")), [])

    def test_renderer_error_is_reported_by_engine(self) -> None:
        with TemporaryDirectory() as temp_dir:
            context = self._context(
                Path(temp_dir),
                renderer=FailingRenderer(RenderingError("renderer failed")),
            )

            result = context.engine.generate(context.request)

            self.assertFalse(result.success)
            self.assertIn("renderer failed", result.error_message)

    def test_pdf_converter_error_is_raised_by_full_pipeline(self) -> None:
        with TemporaryDirectory() as temp_dir:
            context = self._context(Path(temp_dir))

            with patch(
                "src.services.pdf.converter.subprocess.run",
                return_value=subprocess.CompletedProcess(["soffice"], 1, "", "bad"),
            ):
                result = context.engine.generate(context.request)

            self.assertFalse(result.success)
            self.assertIn("LibreOffice conversion failed", result.error_message)

    def test_cleanup_error_is_raised_by_full_pipeline(self) -> None:
        with TemporaryDirectory() as temp_dir:
            context = self._context(
                Path(temp_dir),
                cleanup_service=FailingCleanupService(),
            )

            with self._mock_libreoffice():
                result = context.engine.generate(context.request)

            self.assertTrue(result.success)
            self.assertIsNotNone(result.cleanup_error)

    def test_parser_error_is_reported_by_engine(self) -> None:
        with TemporaryDirectory() as temp_dir:
            context = self._context(Path(temp_dir), parser=FailingParser())

            result = context.engine.generate(context.request)

            self.assertFalse(result.success)
            self.assertIn("parser failed", result.error_message)

    def test_google_error_is_reported_by_engine(self) -> None:
        with TemporaryDirectory() as temp_dir:
            context = self._context(
                Path(temp_dir),
                cache_error=GoogleDriveError("google failed"),
            )

            result = context.engine.generate(context.request)

            self.assertFalse(result.success)
            self.assertIn("google failed", result.error_message)

    def test_unknown_project_is_reported_by_engine(self) -> None:
        with TemporaryDirectory() as temp_dir:
            context = self._context(Path(temp_dir))
            request = ContractRequest("missing", "developer", EMPLOYEE_TEXT)

            result = context.engine.generate(request)

            self.assertFalse(result.success)
            self.assertIn("Project is not configured", result.error_message)

    def test_unknown_vacancy_is_reported_by_engine(self) -> None:
        with TemporaryDirectory() as temp_dir:
            context = self._context(Path(temp_dir))
            request = ContractRequest("project", "missing", EMPLOYEE_TEXT)

            result = context.engine.generate(request)

            self.assertFalse(result.success)
            self.assertIn("Vacancy is not configured", result.error_message)

    def test_benchmark_pipeline_steps(self) -> None:
        with TemporaryDirectory() as temp_dir:
            context = self._context(Path(temp_dir))
            parser = EmployeeParser()
            registry = self._registry()
            renderer = DocxRenderer()
            pdf_converter = context.pdf_converter

            start = time.perf_counter()
            raw_fields = parser.parse(EMPLOYEE_TEXT)
            parser_time = time.perf_counter() - start

            start = time.perf_counter()
            processed_fields = registry.process(raw_fields)
            registry_time = time.perf_counter() - start

            render_output = Path(temp_dir) / "tmp" / "benchmark.docx"
            start = time.perf_counter()
            renderer.render(context.template_path, processed_fields, render_output)
            renderer_time = time.perf_counter() - start

            with self._mock_libreoffice():
                start = time.perf_counter()
                pdf_converter.convert(render_output, Path(temp_dir) / "pdf")
                pdf_time = time.perf_counter() - start

                start = time.perf_counter()
                context.engine.generate(context.request)
                full_time = time.perf_counter() - start

            metrics = {
                "parser": parser_time,
                "processor_registry": registry_time,
                "renderer": renderer_time,
                "pdf": pdf_time,
                "full_generation": full_time,
            }

            self.assertTrue(all(value >= 0 for value in metrics.values()))
            self.assertLess(metrics["full_generation"], 5.0)
            self._write_benchmark_artifact(Path(temp_dir), metrics)

    def _context(
        self,
        root: Path,
        *,
        parser=None,
        renderer=None,
        cleanup_service: CleanupService | None = None,
        cache_error: Exception | None = None,
        with_pdf: bool = True,
    ):
        template_path = root / "template.docx"
        self._create_template(template_path)
        output_docx = root / "tmp" / "contract.docx"
        output_docx.parent.mkdir()
        pdf_dir = root / "pdf"
        pdf_dir.mkdir()
        libreoffice = root / "soffice.exe"
        libreoffice.write_bytes(b"fake executable")

        catalog = TemplateCatalog.from_mapping(
            {
                "projects": [
                    {
                        "id": "project",
                        "name": "Project",
                        "vacancies": [
                            {
                                "id": "developer",
                                "name": "Developer",
                                "template_id": "contract-template",
                            }
                        ],
                    }
                ],
                "templates": [
                    {
                        "id": "contract-template",
                        "name": "Contract",
                        "google_drive_file_id": "drive-file",
                    }
                ],
            }
        )
        pdf_converter = PdfConverter(
            SimpleNamespace(
                libreoffice=SimpleNamespace(
                    executable_path=libreoffice,
                    timeout_seconds=60,
                )
            )
        )
        cleanup = cleanup_service or CleanupService()
        engine = ContractEngine(
            template_catalog=catalog,
            template_cache=LocalTemplateCache(template_path, error=cache_error),
            employee_parser=parser or EmployeeParser(),
            processor_registry=self._registry(),
            docx_renderer=renderer or DocxRenderer(),
            pdf_converter=pdf_converter if with_pdf else None,
            cleanup_service=cleanup if with_pdf else None,
            output_dir=root / "tmp",
            pdf_output_dir=pdf_dir,
        )
        request = ContractRequest("project", "developer", EMPLOYEE_TEXT, output_docx)

        return SimpleNamespace(
            engine=engine,
            request=request,
            template_path=template_path,
            pdf_converter=pdf_converter,
            saved_docx_before_cleanup=template_path,
        )

    def _registry(self) -> ProcessorRegistry:
        registry = ProcessorRegistry()
        registry.register(FioProcessor())
        registry.register(DateProcessor())
        registry.register(MoneyProcessor())
        registry.register(PlaceholderProcessor())
        return registry

    def _create_template(self, path: Path) -> None:
        document = Document()
        first = document.add_paragraph()
        first.add_run("Сотрудник ")
        first.add_run("<ФИО>").bold = True
        document.add_paragraph("<Ф> <И> <О>")
        document.add_paragraph("ИНН: <ИНН>")
        document.add_paragraph("Ставка: <Ставка>")
        document.add_paragraph("Плата: <Плата>")
        document.add_paragraph("Дата выдачи: <Дата выдачи>")
        document.add_paragraph("Дата части: <День> <Месяц> <Год>")
        document.add_paragraph("Повтор: <ФИО> / <ФИО>")

        table = document.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "ФИО"
        table.cell(0, 1).text = "<ФИО>"
        table.cell(1, 0).text = "ИНН"
        table.cell(1, 1).text = "<ИНН>"

        section = document.sections[0]
        section.header.paragraphs[0].text = "Header <ФИО>"
        section.footer.paragraphs[0].text = "Footer <ИНН>"
        document.save(path)

    def _mock_libreoffice(self):
        def run(command, **kwargs):
            output_dir = Path(command[5])
            input_docx = Path(command[6])
            output_dir.mkdir(parents=True, exist_ok=True)
            (output_dir / f"{input_docx.stem}.pdf").write_bytes(self._minimal_pdf())
            return subprocess.CompletedProcess(command, 0, "converted", "")

        return patch("src.services.pdf.converter.subprocess.run", side_effect=run)

    def _minimal_pdf(self) -> bytes:
        objects = [
            b"<< /Type /Catalog /Pages 2 0 R >>",
            b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
            b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 200 200] /Contents 4 0 R >>",
            b"<< /Length 44 >>\nstream\nBT /F1 12 Tf 10 10 Td (Contract) Tj ET\nendstream",
        ]
        content = bytearray(b"%PDF-1.4\n")
        offsets = [0]
        for index, body in enumerate(objects, start=1):
            offsets.append(len(content))
            content.extend(f"{index} 0 obj\n".encode("ascii"))
            content.extend(body)
            content.extend(b"\nendobj\n")
        xref_offset = len(content)
        content.extend(f"xref\n0 {len(objects) + 1}\n".encode("ascii"))
        content.extend(b"0000000000 65535 f \n")
        for offset in offsets[1:]:
            content.extend(f"{offset:010d} 00000 n \n".encode("ascii"))
        content.extend(
            (
                f"trailer\n<< /Size {len(objects) + 1} /Root 1 0 R >>\n"
                f"startxref\n{xref_offset}\n%%EOF\n"
            ).encode("ascii")
        )
        return bytes(content)

    def _assert_pdf_has_pages(self, pdf_path: Path) -> None:
        data = pdf_path.read_bytes()
        self.assertTrue(data.startswith(b"%PDF"))
        self.assertGreater(data.count(b"/Type /Page"), 0)

    def _document_text(self, document: Document) -> str:
        parts = [paragraph.text for paragraph in document.paragraphs]
        for table in document.tables:
            parts.extend(self._table_text(table))
        parts.append(self._headers_text(document))
        parts.append(self._footers_text(document))
        return "\n".join(parts)

    def _table_text(self, table) -> list[str]:
        texts: list[str] = []
        for row in table.rows:
            for cell in row.cells:
                texts.extend(paragraph.text for paragraph in cell.paragraphs)
                for nested_table in cell.tables:
                    texts.extend(self._table_text(nested_table))
        return texts

    def _headers_text(self, document: Document) -> str:
        return "\n".join(
            paragraph.text
            for section in document.sections
            for paragraph in section.header.paragraphs
        )

    def _footers_text(self, document: Document) -> str:
        return "\n".join(
            paragraph.text
            for section in document.sections
            for paragraph in section.footer.paragraphs
        )

    def _write_benchmark_artifact(self, root: Path, metrics: dict[str, float]) -> None:
        (root / "benchmark.json").write_text(
            json.dumps(metrics, ensure_ascii=False, indent=2),
            encoding="utf-8",
        )


class FailingRenderer:
    """Renderer test double that raises a configured error."""

    def __init__(self, error: Exception) -> None:
        self.error = error

    def render(self, template_path, data, output_path) -> RenderResult:
        """Raise the configured renderer error."""
        raise self.error


class FailingParser:
    """Parser test double that fails."""

    def parse(self, text: str) -> dict[str, str]:
        """Raise a parser error."""
        raise ValueError("parser failed")


class FailingCleanupService(CleanupService):
    """Cleanup service test double that fails."""

    def cleanup(self, docx_path, pdf_path, temp_dir, *, delete_docx=True):
        """Raise a cleanup error."""
        raise RuntimeError("cleanup failed")


if __name__ == "__main__":
    unittest.main()
