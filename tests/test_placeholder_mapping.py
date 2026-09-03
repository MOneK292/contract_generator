"""Tests for template placeholder mapping reports."""

from __future__ import annotations

import unittest
from pathlib import Path
from tempfile import TemporaryDirectory

from docx import Document

from src.services.docx.placeholder_mapping import PlaceholderMappingBuilder


class PlaceholderMappingBuilderTest(unittest.TestCase):
    """Placeholder mapping builder behavior."""

    def test_builds_placeholder_map_with_explicit_aliases(self) -> None:
        with TemporaryDirectory() as temp_dir:
            template_path = Path(temp_dir) / "template.docx"
            document = Document()
            document.add_paragraph("<Ф> <И> <О> <ФИО>")
            document.add_paragraph("<Серия, номер П> <Плата> <Регистрация>")
            document.save(template_path)

            report = PlaceholderMappingBuilder().build_for_template(template_path)

            mapping = {match.placeholder: match for match in report.matches}
            self.assertEqual(mapping["Ф"].source, "Ф")
            self.assertFalse(mapping["Ф"].automatic)
            self.assertEqual(mapping["ФИО"].source, "Ф + И + О")
            self.assertTrue(mapping["ФИО"].automatic)
            self.assertEqual(mapping["Серия, номер П"].source, "Серия + Номер П")
            self.assertTrue(mapping["Серия, номер П"].automatic)
            self.assertEqual(mapping["Плата"].source, "Ставка")
            self.assertTrue(mapping["Плата"].automatic)
            self.assertEqual(mapping["Регистрация"].source, "Регистрация")
            self.assertFalse(mapping["Регистрация"].automatic)


if __name__ == "__main__":
    unittest.main()
